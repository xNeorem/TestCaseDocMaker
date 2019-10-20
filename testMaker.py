from docx import Document
from docx.enum.text import WD_LINE_SPACING
import copy

d = {}
with open("dati.txt") as f:
    for line in f:
        (key, *val) = line.split()
        for e in range(0, len(val)):
            if val[e] == "null":
                val[e] = " "
        d[key] = val


lista = d['NeoremID']
rad = d['NeoremID'][0]
lun = int(d['NeoremID'][1])
lista[0] = rad+str(1)
lista[1] = rad+str(2)

for i in range(3, lun+1):
    lista.append(rad + str(i))

for key, value in d.items():
    if(len(value) != lun):
        val = value[len(value)-1]
        for i in range(len(value), lun):
            value.append(val)

for key, value in d.items():
    print(key + " -> " + str(value))

result = []

for i in range(0, lun):
    document = Document("./prova.docx")
    for table in document.tables:

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key in d:
                        if key in paragraph.text:
                            text = paragraph.text.replace(key, d[key][i])
                            style = paragraph.style
                            paragraph.text = text
                            paragraph.style = style

    result.append(document)


def _combine_docx(documents):
    '''
    Takes a list of docx.Documents and performs a deepcopy on the
    first table in each document and adds the copies successively to
    the first document in the list.
   '''

    if not documents or len(documents) == 0:
        return Document()
    if len(documents) == 1:
        return documents[0]

    merged_document = documents[0]
    p = merged_document.paragraphs[0]
    for doc in documents[1:]:
        doc_table = doc.tables[0]

        for i in range(0, 2):  # spacing
            newp = merged_document.add_paragraph("")
            newp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            p._p.addnext(newp._p)

        new_tbl = copy.deepcopy(doc_table._tbl)
        p._p.addnext(new_tbl)

    return merged_document


# a = result[0]
# del result[0]
result.reverse()
# result.insert(0, a)
result = _combine_docx(result)

result.save("result.docx")

print("test case creati con successo in result.docx")

f.close()
